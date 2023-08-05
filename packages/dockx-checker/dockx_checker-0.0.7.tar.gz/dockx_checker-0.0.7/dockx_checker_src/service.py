from flask import *
import os
import tempfile
import shutil
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

STATIC_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'static')
app = Flask(__name__,)

TMP = tempfile.mkdtemp()

KEYS = {}
with open(Path("~").expanduser()/"Desktop"/"hit.txt", "rb") as f:
    buf = f.read()
    lines = []
    try:
        lines = buf.decode("utf-8").split("\n")
    except:
        lines = buf.decode("gbk","ignore").split("\n")

    for l in lines:
        if l.strip() != "" and ":" in l:
            k,v = l.strip().split(":")
            KEYS[k.strip()] = v.strip()

def DealDocx(path):
    doc = Document(path)
    for p in doc.paragraphs:
        for k in KEYS:
            
            if k in p.text:
                old_coms = {}
                old_texts = [r.text for r in p.runs]
                old_style = {r.text: {
                    "style":r.style,
                    "bold":r.font.bold,
                    "italic":r.font.italic,
                    "fontsize":r.font.size,
                    "fontname":r.font.name,
                    "underline":r.font.underline,
                } for r in p.runs}
                ts = []
                for r in p.runs:
                    if len(r.comments)>0:
                        print("save:",r.text, r.comments[0].text)
                        old_coms[r.text] = r.comments[0].text
                    if k in r.text:
                        ts += ("<---->"+k+"<---->").join(r.text.split(k)).split("<---->")
                    else:
                        ts += [r.text]
                p.runs.clear()
                p.text = ""
                for t in ts:
                    r = p.add_run()
                    r.text = t
                    if t == k:
                        r.add_comment(KEYS[k])
                    styles = old_style.get(t,old_style.get(old_texts[0]))
                    r.style = styles["style"]
                    r.bold = styles["bold"]
                    r.italic = styles["italic"]
                    r.font.name = styles["fontname"]
                    if r.element.rPr.rFonts is not None:
                        r.element.rPr.rFonts.set(qn('w:eastAsia'), styles["fontname"])
                        r.font.size = styles["fontsize"]
                        r.underline = styles["underline"]
                        print("------  font add:",t)
                        print(f"style: {r.style} bold: {r.bold} italic: {r.italic} fontname: {r.font.name} fontsize: {r.font.size} underline: {r.underline} ")
                        print("------------------")
                    else:
                        print("------  add:",t)
                        print(f"style: {r.style} bold: {r.bold} italic: {r.italic} fontname: {r.font.name} fontsize: {r.font.size} underline: {r.underline} ")
                        print("------------------")
                    if t in old_coms:
                        r.add_comment(old_coms[t])
                print(k, "->", KEYS[k])
                
                
    new_path =path.replace(".docx","_checked.docx")
    if os.path.exists(new_path):
        os.remove(str(new_path))
    doc.save(path.replace(".docx","_checked.docx"))
    return path.replace(".docx","_checked.docx")

## a service for deal upload file
@app.route('/', methods=['POST','GET'])
def upload():
    if request.method == 'GET':
        # return render_template_string(
        return render_template('index.html')
    else:
        if 'file' not in request.files:
            return 'No file part in the request', 400
        file = request.files['file']
        if file.filename == '':
            return 'No selected file', 400
        if file:
            filename = file.filename
            file.save(os.path.join(TMP, filename)) # 请替换 '/path/to/save' 为你希望保存文件的路径
            filename = DealDocx(os.path.join(TMP, filename))
            dst = os.path.join(STATIC_PATH,"downloads", os.path.basename(filename))
            if os.path.exists(dst):
                os.remove(dst)
            shutil.move(filename, os.path.join(STATIC_PATH,"downloads", os.path.basename(filename)))

            return json.dumps({
                "url":"/static/downloads/"+os.path.basename(filename),
                "filename":os.path.basename(filename),
            }),200
