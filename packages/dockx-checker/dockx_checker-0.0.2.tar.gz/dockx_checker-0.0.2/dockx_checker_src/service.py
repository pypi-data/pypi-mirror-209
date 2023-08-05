from flask import *
import os
from pathlib import Path
from docx import Document
STATIC_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'static')
app = Flask(__name__,)


KEYS = {}
with open(Path("~").expanduser()/"Desktop"/"hit.txt", encoding="gbk") as f:
    for l in f:
        if l.strip() != "":
            k,v = l.strip().split(":")
            KEYS[k.strip()] = v.strip()

def DealDocx(path):
    doc = Document(path)
    for p in doc.paragraphs:
        for k in KEYS:
            
            if k in p.text:
                # p_style = p.style
                style = p.runs[0].style
                bold = p.runs[0].bold
                italic = p.runs[0].italic
                fontsize = p.runs[0].font.size
                fontname = p.runs[0].font.name
                underline = p.runs[0].underline
                ix = p.text.find(k)
                ex = ix + len(k)
                pre = p.text[:ix]
                end = p.text[ex:]
                p.text = pre
                for run in p.runs:
                    run.style = style
                    run.bold = bold
                    run.italic = italic
                    run.font.name = fontname
                    run.font.size = fontsize
                    run.underline = underline
                    
                scrun = p.add_run(k)
                scrun.style = style
                scrun.bold = bold
                scrun.italic = italic
                scrun.font.name = fontname
                scrun.font.size = fontsize
                scrun.underline = underline

                endrun = p.add_run(end)
                endrun.style = style
                endrun.bold = bold
                endrun.italic = italic
                endrun.font.name = fontname
                endrun.font.size = fontsize
                endrun.underline = underline

                print(k, "->", KEYS[k],ix,ex)
                scrun.add_comment(KEYS[k])
    
    doc.save(path.replace(".docx","_checked.docx"))
    return path.replace(".docx","_checked.docx")

## a service for deal upload file
@app.route('/', methods=['POST','GET'])
def upload():
    if request.method == 'GET':
        return render_template('index.html')
    else:
        if 'file' not in request.files:
            return 'No file part in the request', 400
        file = request.files['file']
        if file.filename == '':
            return 'No selected file', 400
        if file:
            filename = file.filename
            file.save(os.path.join('/tmp/', filename)) # 请替换 '/path/to/save' 为你希望保存文件的路径
            filename = DealDocx(os.path.join('/tmp/', filename))
            os.rename(filename, os.path.join(STATIC_PATH,"downloads", os.path.basename(filename)))

            return json.dumps({
                "url":"/static/downloads/"+os.path.basename(filename),
                "filename":os.path.basename(filename),
            }),200
